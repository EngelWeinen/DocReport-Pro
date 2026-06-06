import os
os.environ["PYTHONIOENCODING"] = "utf-8"

import time
import json
import requests
import smtplib
from tkinter import filedialog, messagebox
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
import customtkinter as ctk
from docx import Document

API_URL = "https://open.bigmodel.cn/api/paas/v4/chat/completions"
API_KEY = "your-api-key-here"
MODEL = "glm-4-flash"

SMTP_SERVER = "smtp.qq.com"
SMTP_PORT = 587
SENDER = "your-email@qq.com"
AUTH_CODE = "your-auth-code"
RECEIVER = "receiver@example.com"

ctk.set_appearance_mode("light")
ctk.set_default_color_theme("blue")


class DocReportPro(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("DocReport Pro")
        self.geometry("720x540")
        self.resizable(False, False)
        self.file = None
        self._build_ui()

    def _build_ui(self):
        ctk.CTkLabel(self, text="DocReport Pro", font=("Arial", 28, "bold")).pack(pady=(18, 0))
        ctk.CTkLabel(self, text="by EngelWeinen", font=("Arial", 11, "italic"),
                     text_color=("gray50", "gray70")).pack(pady=(0, 16))

        ctk.CTkButton(self, text="Select Word File", command=self._pick_file,
                      width=240, height=40).pack(pady=4)
        self.status = ctk.CTkLabel(self, text="No file selected", font=("Arial", 12))
        self.status.pack(pady=6)
        ctk.CTkButton(self, text="Generate & Send", command=self._run,
                      width=320, height=48, font=("Arial", 14, "bold")).pack(pady=18)

        self.log = ctk.CTkTextbox(self, width=640, height=180)
        self.log.pack(pady=(4, 12))
        self._log("DocReport Pro initialized")

    def _log(self, msg: str):
        self.log.insert("end", f"[EW] {time.strftime('%H:%M:%S')} {msg}\n")
        self.log.see("end")
        self.update_idletasks()

    def _pick_file(self):
        path = filedialog.askopenfilename(title="Choose .docx", filetypes=[("Word", "*.docx")])
        if not path:
            return
        self.file = path
        self.status.configure(text=f"Ready: {os.path.basename(path)}")
        self._log(f"Loaded {path}")

    def _read_docx(self, path: str) -> str:
        return "\n".join(p.text.strip() for p in Document(path).paragraphs if p.text.strip())

    def _summarize(self, text: str) -> str:
        sys_msg = "你是专业文档分析师，提炼关键信息、数据、结论，生成精简工作报告，篇幅控制在原文1/4，排版分段清晰。"
        headers = {
            "Authorization": f"Bearer {API_KEY}",
            "Content-Type": "application/json"
        }
        payload = {
            "model": MODEL,
            "messages": [
                {"role": "system", "content": sys_msg},
                {"role": "user", "content": f"原文：{text}\n输出精简汇总报告："}
            ],
            "temperature": 0.3,
            "max_tokens": 2500
        }
        resp = requests.post(API_URL, headers=headers, json=payload, timeout=120)
        resp.raise_for_status()
        data = resp.json()
        return data["choices"][0]["message"]["content"].strip()

    def _save_report(self, summary: str) -> str:
        doc = Document()
        doc.add_heading("Document Summary Report", level=1)
        doc.add_paragraph(summary)
        doc.save("AI_Summary_Report.docx")
        return "AI_Summary_Report.docx"

    def _send_mail(self, filepath: str):
        msg = MIMEMultipart()
        msg["From"] = SENDER
        msg["To"] = RECEIVER
        msg["Subject"] = "AI Summary Report"

        with open(filepath, "rb") as f:
            part = MIMEBase("application", "vnd.openxmlformats-officedocument.wordprocessingml.document")
            part.set_payload(f.read())
            encoders.encode_base64(part)
            part.add_header("Content-Disposition", f'attachment; filename="{os.path.basename(filepath)}"')
            msg.attach(part)

        with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as server:
            server.starttls()
            server.login(SENDER, AUTH_CODE)
            server.send_message(msg)

    def _run(self):
        if not self.file:
            messagebox.showwarning("DocReport Pro", "Select a Word file first.")
            return
        try:
            self._log("Parsing document...")
            raw = self._read_docx(self.file)
            self._log("Requesting AI...")
            summary = self._summarize(raw)
            self._log("Building report...")
            report = self._save_report(summary)
            self._log(f"Report saved: {report}")
        except Exception as e:
            self._log(f"AI/Report Error: {e}")
            messagebox.showerror("DocReport Pro", str(e))
            return

        try:
            self._log("Sending email...")
            self._send_mail(report)
            self._log("Email sent.")
            messagebox.showinfo("DocReport Pro", "Report generated & sent.")
        except Exception as e:
            self._log(f"Email Error: {e}")
            messagebox.showwarning("DocReport Pro",
                f"Report saved as {report}, but email failed.\n\n{e}")


if __name__ == "__main__":
    DocReportPro().mainloop()